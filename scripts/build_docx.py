#!/usr/bin/env python3
"""
Genera documento Word profesional consolidado desde archivos markdown.
Corre automáticamente en GitHub Actions en cada push.

Características:
- Parsea markdown correctamente (títulos, listas, tablas, código)
- Aplica estilos profesionales y consistentes
- Genera índice automático
- Maneja citas y bloques de énfasis
- Consolida múltiples archivos en orden
"""

import os
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colores corporativos (ajusta según tu branding)
COLOR_PRIMARY = RGBColor(0, 51, 102)      # Azul oscuro
COLOR_ACCENT = RGBColor(0, 102, 204)     # Azul claro
COLOR_GRAY = RGBColor(96, 96, 96)        # Gris oscuro
COLOR_LIGHT_GRAY = RGBColor(240, 240, 240)  # Gris claro

def add_shading(paragraph, color):
    """Añade fondo a un párrafo."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def parse_markdown_line(line):
    """Identifica el tipo de línea markdown."""
    line = line.rstrip()

    if not line:
        return ('empty', '')

    if line.startswith('# '):
        return ('h1', line[2:].strip())
    elif line.startswith('## '):
        return ('h2', line[3:].strip())
    elif line.startswith('### '):
        return ('h3', line[4:].strip())
    elif line.startswith('#### '):
        return ('h4', line[5:].strip())
    elif line.startswith('##### '):
        return ('h5', line[6:].strip())
    elif line.startswith('> '):
        return ('blockquote', line[2:].strip())
    elif line.startswith('- ') or line.startswith('* '):
        return ('list_item', line[2:].strip())
    elif line.startswith('| '):
        return ('table_row', line)
    elif line.startswith('```'):
        return ('code_fence', line[3:].strip())
    else:
        return ('paragraph', line)

def format_inline_markdown(text):
    """Aplica formato inline: **bold**, *italic*, `code`."""
    # Este es un retorno simplificado, se aplica en add_formatted_paragraph
    return text

def add_formatted_paragraph(doc, text, style='Normal', color=None, size=None):
    """Añade párrafo con soporte para **bold**, *italic*, `code`."""
    para = doc.add_paragraph(style=style)

    # Regex para encontrar patrones markdown
    pattern = r'(\*\*.*?\*\*)|(\*.*?\*)|(`.*?`)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Texto normal antes del match
        if match.start() > last_end:
            run = para.add_run(text[last_end:match.start()])
            if color:
                run.font.color.rgb = color
            if size:
                run.font.size = size

        # Texto formateado
        matched_text = match.group(0)
        if matched_text.startswith('**'):
            content = matched_text[2:-2]
            run = para.add_run(content)
            run.bold = True
            if color:
                run.font.color.rgb = color
            if size:
                run.font.size = size
        elif matched_text.startswith('*'):
            content = matched_text[1:-1]
            run = para.add_run(content)
            run.italic = True
            if color:
                run.font.color.rgb = color
            if size:
                run.font.size = size
        elif matched_text.startswith('`'):
            content = matched_text[1:-1]
            run = para.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            if color is None:
                run.font.color.rgb = RGBColor(192, 0, 0)

        last_end = match.end()

    # Texto restante
    if last_end < len(text):
        run = para.add_run(text[last_end:])
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size

    return para

def parse_table_markdown(lines, start_idx):
    """Parsea una tabla markdown y retorna (tabla_data, siguiente_índice)."""
    table_lines = []
    idx = start_idx

    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break

        # Extrae celdas de la fila
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        table_lines.append(cells)
        idx += 1

    return table_lines, idx

def add_table_from_markdown(doc, table_lines):
    """Crea una tabla en el documento desde datos de tabla markdown."""
    if not table_lines or len(table_lines) < 2:
        return

    # La segunda fila en markdown es separador, la saltamos
    headers = table_lines[0]
    data_rows = table_lines[2:] if len(table_lines) > 2 else []

    # Crea tabla con headers
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Formato de headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Estilo del header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Fondo del header (shading en la celda, no en párrafo)
        try:
            tcPr = header_cells[i]._element.get_or_add_tcPr()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '003366')
            tcPr.append(shading_elm)
        except:
            pass  # Si falla, continúa sin shading

    # Datos
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = cell_data

def read_markdown_file(file_path):
    """Lee un archivo markdown."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"⚠️  Archivo no encontrado: {file_path}")
        return None

def add_cover_page(doc):
    """Añade portada profesional."""
    # Título principal
    title = doc.add_heading('Estado en el Territorio', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_format = title.runs[0].font
    title_format.size = Pt(28)
    title_format.color.rgb = COLOR_PRIMARY
    title_format.bold = True

    doc.add_paragraph()  # Espacio

    # Subtítulo
    subtitle = doc.add_paragraph('Plan de Gobierno Integral 2026-2030')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(16)
    subtitle_format.italic = True
    subtitle_format.color.rgb = COLOR_ACCENT

    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio

    # Descripción
    desc = doc.add_paragraph(
        'Propuesta para llevar el Estado completo, coordinado y verificable '
        'a territorios priorizados, concentrando inversión hasta que cambien, '
        'expandiendo por resultado y no por calendario.'
    )
    desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in desc.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_GRAY

    # Espacios
    for _ in range(6):
        doc.add_paragraph()

    # Fecha y generación
    fecha = doc.add_paragraph(f'Generado: {datetime.now().strftime("%d de %B de %Y")}')
    fecha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fecha.runs[0].font.size = Pt(10)
    fecha.runs[0].font.italic = True
    fecha.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

def add_table_of_contents(doc, sections):
    """Añade tabla de contenidos."""
    doc.add_heading('Contenido', 1)

    for i, (_, section_name) in enumerate(sections, 1):
        para = doc.add_paragraph(section_name, style='List Number')
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_page_break()

def process_markdown_file(doc, file_path, section_title):
    """Procesa un archivo markdown y lo añade al documento."""
    content = read_markdown_file(file_path)
    if not content:
        return False

    lines = content.split('\n')
    idx = 0
    in_code_block = False
    code_language = ''
    code_lines = []

    # Salta líneas de encabezado o comentarios iniciales
    while idx < len(lines) and lines[idx].startswith('>'):
        idx += 1

    while idx < len(lines):
        line = lines[idx]
        line_type, content_text = parse_markdown_line(line)

        # Manejo de bloques de código
        if line_type == 'code_fence':
            if not in_code_block:
                in_code_block = True
                code_language = content_text
                code_lines = []
            else:
                # Cierra bloque de código
                in_code_block = False
                code_para = doc.add_paragraph()
                for code_line in code_lines:
                    code_para.add_run(code_line + '\n').font.name = 'Courier New'
                    code_para.runs[-1].font.size = Pt(8)
                add_shading(code_para, 'F5F5F5')
            idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        # Procesa línea normal
        if line_type == 'empty':
            idx += 1
        elif line_type == 'h1':
            heading = doc.add_heading(content_text, level=1)
            heading.runs[0].font.color.rgb = COLOR_PRIMARY
            idx += 1
        elif line_type == 'h2':
            heading = doc.add_heading(content_text, level=2)
            heading.runs[0].font.color.rgb = COLOR_ACCENT
            idx += 1
        elif line_type == 'h3':
            heading = doc.add_heading(content_text, level=3)
            idx += 1
        elif line_type == 'h4':
            heading = doc.add_heading(content_text, level=4)
            idx += 1
        elif line_type == 'blockquote':
            # Cita destacada
            para = doc.add_paragraph()
            add_shading(para, 'F0F0F0')
            run = para.add_run(content_text)
            run.italic = True
            run.font.color.rgb = COLOR_GRAY
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)
            idx += 1
        elif line_type == 'list_item':
            doc.add_paragraph(content_text, style='List Bullet')
            idx += 1
        elif line_type == 'table_row':
            table_lines, next_idx = parse_table_markdown(lines, idx)
            if table_lines:
                add_table_from_markdown(doc, table_lines)
            idx = next_idx
        elif line_type == 'paragraph':
            if content_text:
                add_formatted_paragraph(doc, content_text)
            idx += 1
        else:
            idx += 1

    return True

def build_document():
    """Construye el documento Word completo."""

    print("🔨 Construyendo documento Word...")
    doc = Document()

    # Márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Portada
    add_cover_page(doc)

    # Archivos a incluir (en orden)
    # Solo incluye los que sabemos que existen
    files_to_include = [
        ('docs/05-columna-vertebral.md', 'Columna Vertebral'),
        ('docs/frentes/frente-02-reforma-agraria.md', 'Frente II: Reforma Agraria'),
        ('docs/frentes/frente-08-capacidad-estatal.md', 'Frente VIII: Capacidad Operativa'),
        ('docs/frentes/frente-11-servicios-basicos.md', 'Frente XI: Servicios Básicos'),
        ('docs/frentes/frente-13-integridad-mecanismos.md', 'Frente XIII: Integridad Estructural'),
        ('docs/transversal/fundamentos-investigacion.md', 'Fundamentos en Investigación Reciente'),
        ('docs/transversal/dimension-deliberativa.md', 'Dimensión Deliberativa'),
    ]

    # Verifica cuáles existen realmente
    valid_files = []
    for file_path, section_name in files_to_include:
        if os.path.exists(file_path):
            valid_files.append((file_path, section_name))

    if not valid_files:
        print("  ❌ No se encontraron archivos markdown para procesar")
        return False

    # Tabla de contenidos
    add_table_of_contents(doc, valid_files)

    # Procesa cada archivo válido
    for file_path, section_name in valid_files:
        try:
            print(f"  ✓ {section_name}")
            success = process_markdown_file(doc, file_path, section_name)
            if success:
                doc.add_page_break()
        except Exception as e:
            print(f"  ⚠️  Error procesando {section_name}: {str(e)}")
            continue

    # Pie de página con número de página
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = (
        f"Estado en el Territorio • "
        f"Generado {datetime.now().strftime('%Y-%m-%d %H:%M')} • "
        f"Página "
    )

    # Añade número de página
    run = footer_para.add_run()
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve">PAGE</w:instrText>'.format(nsdecls('w')))
    fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    footer_para.runs[0].font.size = Pt(8)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Guarda
    output_dir = 'dist'
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, 'Estado-en-el-Territorio.docx')
    doc.save(output_path)

    file_size_kb = os.path.getsize(output_path) / 1024
    print(f"\n✅ Documento generado exitosamente")
    print(f"   📄 {output_path}")
    print(f"   📊 Tamaño: {file_size_kb:.1f} KB")
    print(f"   📅 {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

if __name__ == '__main__':
    try:
        build_document()
    except Exception as e:
        print(f"\n❌ Error al generar documento:")
        print(f"   {str(e)}")
        import traceback
        traceback.print_exc()
        exit(1)
                                                                                                                                                                                                                                                   